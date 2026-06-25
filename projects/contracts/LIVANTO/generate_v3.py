# -*- coding: utf-8 -*-
"""Generate Bevollmächtigungsvertrag V3 DOCX using python-docx."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# Page setup
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.space_before = Pt(0)

# Colors
DARK_BLUE = RGBColor(0x1A, 0x3A, 0x5C)
RED = RGBColor(0xAA, 0x00, 0x00)
GRAY = RGBColor(0x55, 0x55, 0x55)
LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)
HEADER_BG = "D6E4F0"


def add_de(de_text, cn_text=None, bold=False, size=10.5, color=None, cn_size=9.5):
    """Add bilingual paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    run_de = p.add_run(de_text)
    run_de.font.name = 'Arial'
    run_de.font.size = Pt(size)
    run_de.bold = bold
    if color:
        run_de.font.color.rgb = color
    if cn_text:
        run_br = p.add_run('\n')
        run_cn = p.add_run(cn_text)
        run_cn.font.name = 'Microsoft YaHei'
        run_cn.font.size = Pt(cn_size)
        run_cn.font.color.rgb = GRAY
    return p


def add_heading_styled(text, level=1):
    """Add a section heading with bottom border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(13) if level == 1 else Pt(11)
    run.bold = True
    run.font.color.rgb = DARK_BLUE
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="4" w:color="1A3A5C"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def add_clause(num_text, de, cn, indent=0):
    """Add a numbered clause with DE + CN."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    if indent > 0:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(num_text + ' ' + de)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run2 = p.add_run('\n' + cn)
    run2.font.name = 'Microsoft YaHei'
    run2.font.size = Pt(9)
    run2.font.color.rgb = GRAY
    return p


def add_sub(a, de, cn):
    """Add sub-clause like a) b) etc."""
    return add_clause(a + ')', de, cn, indent=0.5)


def add_changed_clause(num_text, de, cn, indent=0):
    """Add a CHANGED clause with red marker."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    if indent > 0:
        p.paragraph_format.left_indent = Cm(indent)
    # Red marker
    r0 = p.add_run('[GEÄNDERT] ')
    r0.font.name = 'Arial'
    r0.font.size = Pt(8)
    r0.font.color.rgb = RED
    r0.bold = True
    r1 = p.add_run(num_text + ' ' + de)
    r1.font.name = 'Arial'
    r1.font.size = Pt(10)
    # Chinese with red marker
    r2 = p.add_run('\n[已修改] ')
    r2.font.name = 'Microsoft YaHei'
    r2.font.size = Pt(7)
    r2.font.color.rgb = RED
    r3 = p.add_run(cn)
    r3.font.name = 'Microsoft YaHei'
    r3.font.size = Pt(9)
    r3.font.color.rgb = GRAY
    return p


def add_new_clause(num_text, de, cn):
    """Add a NEW clause with red [NEU] marker."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r0 = p.add_run(num_text + ' ')
    r0.font.name = 'Arial'
    r0.font.size = Pt(10)
    r0.bold = True
    r0.font.color.rgb = RED
    r1 = p.add_run('[NEU] ')
    r1.font.name = 'Arial'
    r1.font.size = Pt(8)
    r1.font.color.rgb = RED
    r1.bold = True
    r2 = p.add_run(de)
    r2.font.name = 'Arial'
    r2.font.size = Pt(10)
    r3 = p.add_run('\n[新增] ')
    r3.font.name = 'Microsoft YaHei'
    r3.font.size = Pt(7)
    r3.font.color.rgb = RED
    r4 = p.add_run(cn)
    r4.font.name = 'Microsoft YaHei'
    r4.font.size = Pt(9)
    r4.font.color.rgb = GRAY
    return p


def add_simple(text, bold=False, size=10.5):
    """Add a simple paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.bold = bold
    return p


def add_page_break():
    doc.add_page_break()


def set_cell_shading(cell, color):
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_signature_table():
    """Add 3-column signature table."""
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    headers = [
        ("Der Kunde\n客户", "LIVANTO GmbH\nDer Bevollmächtigte / 授权代表", "FREDDY (Shanghai)\nZahlungsdienstleister / 支付服务商")
    ]
    for col, text in enumerate(headers[0]):
        cell = table.cell(0, col)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.bold = True

    lines = [
        "_________________________________\nOrt, Datum / 地点, 日期",
        "_________________________________\nUnterschrift / 签名",
        "_________________________________\nName in Druckbuchstaben / 正楷姓名"
    ]
    for row_idx, line in enumerate(lines):
        for col in range(3):
            cell = table.cell(row_idx + 1, col)
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            run = p.add_run(line)
            run.font.name = 'Arial'
            run.font.size = Pt(8.5)
            run.font.color.rgb = LIGHT_GRAY

    return table


# ============================================================
# BUILD DOCUMENT
# ============================================================

# --- TITLE ---
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(30)
p.paragraph_format.space_after = Pt(4)
run = p.add_run('Bevollmächtigungsvertrag')
run.font.name = 'Arial'
run.font.size = Pt(20)
run.bold = True
run.font.color.rgb = DARK_BLUE

add_simple('授权代表合同', bold=True, size=18)

add_simple('gemäß § 35 Absatz 2 Verpackungsgesetz (VerpackG) / PPWR (EU) 2025/40', size=10)
add_simple('依据《德国包装法》§35条第2款 / 欧盟PPWR 2025/40', size=9)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
run = p.add_run('VERSION 3.0 | Stand: Juni 2026 | 版本 3.0 | 日期：2026年6月')
run.font.name = 'Arial'
run.font.size = Pt(8.5)
run.bold = True
run.font.color.rgb = RED

# --- PREAMBLE ---
add_de(
    'Dieser Bevollmächtigungsvertrag ist ein Geschäftsbesorgungsvertrag gemäß § 675 BGB '
    'zur Bestellung eines Bevollmächtigten nach § 35 Absatz 2 VerpackG. Der Bevollmächtigte '
    'handelt im eigenen Namen und gilt als „Quasi-Hersteller“ im Sinne des '
    'Verpackungsgesetzes, ohne selbst Verpackungen in Verkehr zu bringen.',
    '本合同为依据《德国民法典》(BGB)第675条之事务处理合同，旨在根据《德国包装法》(VerpackG)'
    '第35条第2款指定授权代表。授权代表以其自身名义行事，在《包装法》意义上被视为“准生产商”，'
    '但其自身并不投放包装。'
)

add_de(
    'Ab dem 12. August 2026 tritt die EU-Verpackungsverordnung (PPWR) 2025/40 vollständig '
    'in Kraft. Das VerpackG wird durch das Verpackungsrecht-Durchführungsgesetz (VerpackDG) '
    'abgelöst. Die Bestellung eines Bevollmächtigten wird für alle Hersteller ohne '
    'Niederlassung in Deutschland verpflichtend.',
    '自2026年8月12日起，《欧盟包装与包装废弃物法规》(PPWR) 2025/40全面生效，《包装法》将被'
    '《包装法实施法》(VerpackDG)取代。届时所有在德国无分支机构的生产商必须指定授权代表。'
)

# --- §1 ---
add_heading_styled('§1  Vertragsparteien  |  合同双方')

add_simple('Auftraggeber / Kunde:  |  委托方 / 客户：', bold=True, size=11)
add_simple('[Firmenname des Kunden — wird bei Vertragsschluss ausgefüllt]  |  [客户公司名称 — 签约时填写]')
add_simple('Anschrift: [Adresse]  |  地址：[地址]')
add_simple('USt-ID / VAT: [falls vorhanden]  |  增值税号：[如有]')
add_simple('E-Mail: [E-Mail-Adresse]  |  电子邮箱：[邮箱地址]')
add_simple('- nachfolgend „Kunde“ genannt -  |  - 以下称“客户” -')

add_simple('')
add_simple('Bevollmächtigter:  |  授权代表：', bold=True, size=11)
add_simple('LIVANTO GmbH  |  LIVANTO 有限责任公司')
add_simple('Anschrift: Luisenhoffnung 3C, 44225 Dortmund, Deutschland  |  地址：Luisenhoffnung 3C, 44225 Dortmund, Germany')
add_simple('Geschäftsführer: Zifeng Qian  |  总经理：钱子风')
add_simple('Handelsregister: [HRB-Nr.]  |  USt-ID: [DE...]  |  商业登记号：[HRB-Nr.]  |  增值税号：[DE...]')
add_simple('- nachfolgend „Bevollmächtigter“ genannt -  |  - 以下称“授权代表” -')

# --- §2 ---
add_heading_styled('§2  Vertragsgegenstand  |  合同标的')

add_clause('(1)',
    'Der Kunde bestellt hiermit unwiderruflich die LIVANTO GmbH zu seinem Bevollmächtigten '
    'gemäß § 35 Absatz 2 VerpackG. Der Bevollmächtigte übernimmt sämtliche Pflichten des '
    'Kunden als Hersteller nach dem Verpackungsgesetz — mit alleiniger Ausnahme der '
    'Registrierungspflicht nach § 9 VerpackG, die beim Kunden verbleibt. Eine Beschränkung '
    'auf einzelne Pflichten ist ausgeschlossen („Ganz-oder-Gar-Nicht-Lösung“).',
    '(1) 客户在此不可撤销地指定LIVANTO有限责任公司为其依据《包装法》第35条第2款的授权代表。'
    '授权代表承担客户作为生产商在《包装法》下的全部义务——唯一例外为第9条规定的注册义务'
    '（由客户自行履行）。禁止将授权限制于部分义务（“全有或全无原则”）。')

add_clause('(2)',
    'Der Bevollmächtigte handelt im eigenen Namen. Er wird als „Quasi-Hersteller“ '
    'behandelt und ist unmittelbar gegenüber der ZSVR, den dualen Systemen und den '
    'zuständigen Behörden verantwortlich.',
    '(2) 授权代表以其自身名义行事，被视为“准生产商”，直接对包装品中央登记处(ZSVR)、'
    '双元系统和各主管部门负责。')

add_clause('(3)',
    'Der Umfang der Bevollmächtigung umfasst insbesondere die folgenden Pflichten '
    '(nicht abschließende Aufzählung):',
    '(3) 授权范围特别包括以下义务（非穷尽列举）：')

add_sub('a',
    'Abschluss und Aufrechterhaltung des Systembeteiligungsvertrages mit einem dualen System (§ 7 VerpackG)',
    'a) 签订并维持与双元系统的系统参与合同（《包装法》第7条）')
add_sub('b',
    'Datenmeldung an die ZSVR nach Materialfraktionen (§ 10 VerpackG)',
    'b) 按材料分类向ZSVR进行数据申报（《包装法》第10条）')
add_sub('c',
    'Vollständigkeitserklärung, soweit gesetzlich erforderlich (§ 11 VerpackG)',
    'c) 完整性声明（如法律要求，《包装法》第11条）')
add_sub('d',
    'Rücknahmeorganisation für Transport-, Um- und Verkaufsverpackungen (§ 15 VerpackG)',
    'd) 运输包装、外包装和销售包装的回收组织（《包装法》第15条）')
add_sub('e',
    'Pfandbeteiligung für Einweggetränkeverpackungen (§ 31 VerpackG)',
    'e) 一次性饮料包装的押金系统参与（《包装法》第31条）')
add_sub('f',
    'Entgegennahme von behördlichen Mitteilungen und Bescheiden',
    'f) 接收主管部门的通知和决定')

add_clause('(4)',
    'Der Bevollmächtigte ist berechtigt, zur Erfüllung seiner Pflichten Dritte '
    '(z. B. duale Systembetreiber, Wirtschaftsprüfer) einzuschalten.',
    '(4) 授权代表有权为履行其义务而聘请第三方（如双元系统运营商、审计师等）。')

add_clause('(5)',
    'Die vom Kunden angegebenen Verpackungsarten und -mengen (Anlage B) stellen die '
    'bestmögliche Schätzung des Kunden dar. Die tatsächlichen gesetzlichen Pflichten '
    'können darüber hinausgehen.',
    '(5) 客户声明的包装种类和数量（附件B）为其最佳估计。实际法律义务可能超出该范围。')

add_clause('(6)',
    'Eine Lizenzierung von Markenrechten (z. B. „Der Grüne Punkt“) ist nicht '
    'Gegenstand dieses Vertrages.',
    '(6) 商标使用许可（如“绿点”标识）不属于本合同标的。')

add_clause('(7)',
    'Die Bestellung gilt ab Vertragsbeginn für sämtliche vom Kunden im Geltungsbereich '
    'des VerpackG in Verkehr gebrachten Verpackungen. Bei Inkrafttreten des VerpackDG/PPWR '
    'zum 12. August 2026 setzt sich die Bevollmächtigung unter dem neuen Rechtsrahmen fort.',
    '(7) 本授权自合同生效起适用于客户在《包装法》适用范围内投放的所有包装。2026年8月12日'
    '《包装法实施法》/PPWR生效后，授权在新的法律框架下继续有效。')

add_page_break()

# --- §3 ---
add_heading_styled('§3  Pflichten des Kunden  |  客户义务')

add_clause('(1)',
    'Der Kunde ist verpflichtet, sich eigenständig im Verpackungsregister LUCID zu '
    'registrieren (§ 9 VerpackG) und seine Registrierungsdaten stets aktuell zu halten. '
    'Diese Pflicht kann nicht auf den Bevollmächtigten übertragen werden.',
    '(1) 客户有义务自行在LUCID包装品登记处注册（《包装法》第9条），并始终保持注册数据为最新。'
    '该义务不得转移给授权代表。')

add_clause('(2)',
    'Unverzüglich nach Vertragsschluss benennt der Kunde die LIVANTO GmbH in LUCID '
    'als Bevollmächtigten und hält diese Benennung während der gesamten Vertragslaufzeit aufrecht.',
    '(2) 合同签订后，客户应立即在LUCID中指定LIVANTO有限责任公司为授权代表，并在整个合同期内保持该指定。')

add_clause('(3)',
    'Der Kunde übermittelt dem Bevollmächtigten jährlich bis zum 15. Februar des '
    'Folgejahres vollständige und wahrheitsgemäße Angaben zu allen im vorangegangenen '
    'Kalenderjahr in Verkehr gebrachten Verpackungen, aufgeschlüsselt nach '
    'Materialfraktionen und Verpackungskategorien.',
    '(3) 客户每年应在次年2月15日前向授权代表提供上一日历年投放市场的所有包装的完整真实数据，'
    '按材料类别和包装类别细分。')

# CHANGED §3(4)
add_changed_clause('(4)',
    'Änderungen der Verpackungsarten oder eine Abweichung von mehr als 30 % gegenüber '
    'der Erstdeklaration sind dem Bevollmächtigten unverzüglich mitzuteilen. Unterlässt '
    'der Kunde diese Mitteilung schuldhaft, haftet er für sämtliche dem Bevollmächtigten '
    'daraus entstehenden Mehrkosten, Bußgelder oder sonstigen Nachteile.',
    '(4) 包装种类变更或实际数量与首次申报偏差超过30%的，应立即通知授权代表。如客户因过错'
    '未履行该通知义务，客户应对由此给授权代表造成的所有额外费用、罚款或其他损失承担责任。')

add_clause('(5)',
    'Der Kunde bestätigt, seinen Herstellerstatus gemäß VerpackG geprüft zu haben. '
    'Das Risiko einer Fehlklassifizierung trägt der Kunde. Der Kunde benennt einen '
    'Ansprechpartner und stellt sicher, dass Anfragen des Bevollmächtigten innerhalb '
    'von 14 Tagen beantwortet werden.',
    '(5) 客户确认已核实其根据《包装法》的生产商身份，错误分类的风险由客户承担。客户指定一名联系人，'
    '确保在14天内回复授权代表的询问。')

# CHANGED §3(6) — Freistellung with 30-day buffer
add_changed_clause('(6)',
    'Der Kunde stellt den Bevollmächtigten von sämtlichen Ansprüchen Dritter '
    '(ZSVR, duale Systeme, Behörden, Wettbewerber) frei, die aus unrichtigen, '
    'unvollständigen oder verspäteten Angaben des Kunden entstehen. Dies umfasst '
    'ausdrücklich Bußgelder, Sanktionen, Vertriebsverbote, Abmahnkosten und '
    'Rechtsverfolgungskosten. Der Bevollmächtigte kann die Zahlung der entsprechenden '
    'Beträge vom Kunden verlangen. Erfolgt die Zahlung nicht innerhalb von 30 Tagen '
    'nach Zugang der schriftlichen Zahlungsaufforderung, ist der Bevollmächtigte zur '
    'außerordentlichen fristlosen Kündigung berechtigt; in diesem Fall werden bereits '
    'gezahlte Gebühren nicht erstattet. Die Geltendmachung weitergehender '
    'Schadensersatzansprüche bleibt vorbehalten.',
    '(6) 客户应赔偿授权代表因客户提供的不正确、不完整或延迟数据而产生的所有第三方'
    '（包括ZSVR、双元系统、主管部门、竞争对手）索赔，明确包括罚款、制裁、销售禁令、'
    '警告费用和法律追诉费用。授权代表可要求客户支付相应款项。如客户在收到书面付款要求后'
    '30天内未支付，授权代表有权特别无通知终止合同；此种情况下，已支付的费用不予退还。'
    '主张进一步损害赔偿的权利不受影响。')

# --- §4 ---
add_heading_styled('§4  Pflichten des Bevollmächtigten  |  授权代表义务')

add_clause('(1)',
    'Der Bevollmächtigte bestätigt unverzüglich nach Benennung durch den Kunden '
    'die Autorisierung im LUCID-Portal.',
    '(1) 授权代表在客户指定后立即在LUCID门户中确认授权。')

add_clause('(2)',
    'Der Bevollmächtigte schließt einen Systembeteiligungsvertrag mit einem dualen '
    'System ab und hält diesen während der gesamten Vertragslaufzeit aufrecht.',
    '(2) 授权代表与双元系统签订系统参与合同并在整个合同期内维持其有效。')

add_clause('(3)',
    'Der Bevollmächtigte übermittelt die jährlichen Datenmeldungen fristgerecht an '
    'die ZSVR und stellt dem Kunden auf Wunsch Auszüge der übermittelten Daten zur Verfügung.',
    '(3) 授权代表按期向ZSVR提交年度数据申报，并应客户要求提供已提交数据的摘要。')

add_clause('(4)',
    'Sofern die gesetzlichen Schwellenwerte überschritten werden, beauftragt der '
    'Bevollmächtigte einen registrierten Sachverständigen (§ 27 VerpackG) und reicht '
    'die Vollständigkeitserklärung bei der ZSVR ein.',
    '(4) 如超过法定阈值，授权代表聘请注册审计师（《包装法》第27条）并向ZSVR提交完整性声明。')

add_clause('(5)',
    'Der Bevollmächtigte organisiert die Rücknahme und Verwertung für Verpackungen '
    'nach § 15 VerpackG, sofern der Kunde derartige Verpackungen in Verkehr bringt.',
    '(5) 如客户投放《包装法》第15条所指的包装，授权代表组织此类包装的回收和资源化利用。')

add_clause('(6)',
    'Der Bevollmächtigte informiert den Kunden über wesentliche Änderungen der '
    'gesetzlichen Rahmenbedingungen (VerpackG, PPWR, VerpackDG), die dessen Pflichten betreffen.',
    '(6) 授权代表就涉及客户义务的法律框架（《包装法》、PPWR、《包装法实施法》）的重大变更通知客户。')

# --- §5 ---
add_heading_styled('§5  Verpackungsdaten und Meldewesen  |  包装数据与报告')

add_clause('(1)',
    'Bei Vertragsschluss gibt der Kunde seine voraussichtlichen jährlichen '
    'Verpackungsmengen an (Erstdeklaration, Anlage B). Diese dienen als Grundlage '
    'für den Abschluss des dualen Systemvertrages.',
    '(1) 合同签订时，客户提供预计年度包装量（首次申报，附件B），作为签订双元系统合同的基础。')

# CHANGED §5(2) — "ausschließlich" removed
add_changed_clause('(2)',
    'Die jährliche Ist-Mengen-Meldung erfolgt bis zum 15. Februar des Folgejahres '
    'über das Online-Portal des Bevollmächtigten.',
    '(2) 年度实际数量申报截止日期为次年2月15日，通过授权代表的在线门户提交。')

add_clause('(3)',
    'Überschreiten die tatsächlichen Mengen die gemeldeten Mengen um mehr als 20 %, '
    'wird auf die Überschreitungsmenge ein Zuschlag von 20 % erhoben. Unterschreiten '
    'die tatsächlichen Mengen die gemeldeten Mengen, erfolgt eine Rückerstattung nur '
    'bis zu einer Abweichung von 10 %.',
    '(3) 实际数量超出申报数量20%以上的，超出部分加收20%附加费。实际数量低于申报数量的，'
    '仅退还不超过10%的差额。')

add_clause('(4)',
    'Erfolgt keine fristgerechte Meldung durch den Kunden, darf der Bevollmächtigte '
    'die Vorjahresmengen oder — im ersten Jahr — die geschätzten Mengen als '
    'Ist-Mengen behandeln. Eine Erstattung findet in diesem Fall nicht statt.',
    '(4) 客户未按期申报的，授权代表可将上年数量（首年则为预估量）视为实际数量处理。此情形下不予退款。')

# CHANGED §5(5) — technical failure fallback
add_changed_clause('(5)',
    'Meldungen, die ohne vorherige schriftliche Zustimmung des Bevollmächtigten '
    'nicht über das Online-Portal eingereicht werden, gelten als nicht vorgelegt. '
    'Bei nachweislichem technischen Ausfall des Portals kann der Bevollmächtigte '
    'die Einreichung per E-Mail gestatten. Die Zulässigkeit der E-Mail-Einreichung '
    'ist vom Bevollmächtigten vorab schriftlich zu bestätigen.',
    '(5) 未经授权代表事先书面同意，未通过在线门户提交的申报视为未提交。如门户出现可证明的'
    '技术故障，授权代表可允许通过电子邮件提交。电子邮件提交的有效性需由授权代表事先书面确认。')

# --- §6 ---
add_heading_styled('§6  Vergütung und Zahlungsbedingungen  |  费用与支付')

add_clause('(1)',
    'Der Kunde entrichtet eine jährliche Grundgebühr gemäß der gewählten Servicestufe nach Anlage C.',
    '(1) 客户根据附件C所选服务等级支付年度基本费用。')

add_clause('(2)',
    'Zusätzlich zur Grundgebühr trägt der Kunde sämtliche an die dualen Systeme zu '
    'entrichtenden Lizenzentgelte (Systembeteiligungsgebühren). Diese werden in '
    'tatsächlicher Höhe ohne Aufschlag an den Kunden weiterbelastet.',
    '(2) 除基本费用外，客户承担应向双元系统支付的全部许可费用（系统参与费），按实际金额原价过账，不加价。')

add_simple('Servicestufen (siehe Anlage C):  |  服务等级（详见附件C）：')
add_simple('  •  Basis (EUR 89,00/Jahr): Kern-AR-Dienstleistungen  |  基础 (89.00欧元/年)：核心AR服务')
add_simple('  •  Standard (EUR 159,00/Jahr): Basis + prioritäre Reaktionszeit (48 h) + erweiterte Klassifizierungsberatung + jährlicher Compliance-Kurzbericht')
add_simple('      标准 (159.00欧元/年)：基础+优先响应(48h)+扩展分类咨询+年度合规简报')
add_simple('  •  Premium (EUR 249,00/Jahr): Standard + Vollständigkeitserklärungs-Koordination + persönlicher Ansprechpartner + 24h-Reaktionszeit')
add_simple('      高级 (249.00欧元/年)：标准+完整性声明协调+专属客户经理+24h响应')

add_clause('(3)',
    'Alle Beträge verstehen sich netto zuzüglich der gesetzlichen Mehrwertsteuer.',
    '(3) 所有金额均为净价，另加法定的增值税。')

add_clause('(4)',
    'Die Jahresgebühr wird jährlich im Voraus in Rechnung gestellt. Die Zahlung ist '
    'innerhalb von 14 Tagen nach Rechnungsstellung fällig.',
    '(4) 年费每年提前开具发票，付款期限为发票日期后14天。')

# CHANGED §6(5) — "ausschließlich" removed
add_changed_clause('(5)',
    'Die Zahlung erfolgt über den Zahlungsdienstleister FREDDY gemäß Anlage A. '
    'Die Zahlung an FREDDY gilt als wirksame und schuldbefreiende Zahlung an LIVANTO.',
    '(5) 付款通过支付服务商FREDDY（附件A）进行。向FREDDY付款视为向LIVANTO有效且解除债务的付款。')

# CHANGED §6(6) — VPI + regulatory exception
add_changed_clause('(6)',
    'Der Bevollmächtigte ist berechtigt, die Jahresgrundgebühr einmal pro Kalenderjahr '
    'mit einer Frist von zwei Monaten anzupassen. Die Anpassung ist auf die Veränderung '
    'des Verbraucherpreisindex (VPI) für Deutschland beschränkt. Übersteigt die Anpassung '
    '10 %, steht dem Kunden ein Sonderkündigungsrecht mit einer Frist von einem Monat '
    'zum Anpassungszeitpunkt zu. Führen wesentliche Änderungen der gesetzlichen '
    'Rahmenbedingungen (insbesondere VerpackG, PPWR, VerpackDG) zu einer erheblichen '
    'Erhöhung der Betriebskosten des Bevollmächtigten, ist der Bevollmächtigte berechtigt, '
    'die Jahresgrundgebühr über die VPI-Veränderung hinaus angemessen anzupassen; die '
    'Anpassung ist dem Kunden unter Darlegung der erhöhten Kosten schriftlich zu begründen.',
    '(6) 授权代表有权每个日历年调整一次年度基本费用，提前两个月通知。调整幅度以德国消费者价格指数(VPI)'
    '变动为限。如调整超过10%，客户有权在调整生效时一个月内特别终止合同。如法律框架（特别是《包装法》、'
    'PPWR、《包装法实施法》）的重大变更导致授权代表运营成本显著增加，授权代表有权超出VPI变动幅度合理'
    '调整年度基本费用；调整应书面通知客户并说明成本增加的依据。')

add_clause('(7)',
    'Bei Zahlungsverzug gelten die gesetzlichen Verzugszinsen gemäß § 288 BGB.',
    '(7) 逾期付款适用《德国民法典》第288条的法定延迟利息。')

add_page_break()

# --- §7 ---
add_heading_styled('§7  Vollständigkeitserklärung  |  完整性声明')

add_clause('(1)',
    'Die gesetzliche Pflicht zur Abgabe einer Vollständigkeitserklärung (§ 11 VerpackG) '
    'entsteht i. d. R. erst bei Überschreiten der Schwellenwerte von 40.000 kg '
    'systembeteiligungspflichtigen Verpackungen oder 120.000 kg Transport- und '
    'Gewerbeverpackungen pro Kalenderjahr. Der Bevollmächtigte informiert den Kunden, '
    'sofern diese Schwellenwerte erreicht werden.',
    '(1) 完整性声明的法定义务（《包装法》第11条）通常仅在超过系统参与包装40,000公斤/年或'
    '运输/商业包装120,000公斤/年的阈值时产生。如达到上述阈值，授权代表将通知客户。')

add_clause('(2)',
    'Wird die Vollständigkeitserklärung erforderlich, beauftragt der Bevollmächtigte '
    'einen registrierten Sachverständigen (§ 27 VerpackG). Die Kosten hierfür werden in '
    'tatsächlicher Höhe zuzüglich einer Bearbeitungsgebühr von EUR 50,00 an den Kunden '
    'weiterbelastet.',
    '(2) 如需提交完整性声明，授权代表聘请注册审计师（《包装法》第27条），相关费用按实际金额'
    '加50.00欧元手续费向客户收取。')

add_clause('(3)',
    'Der Kunde verpflichtet sich, dem Bevollmächtigten und dem Sachverständigen die '
    'erforderlichen Unterlagen innerhalb von 14 Tagen nach Aufforderung zur Verfügung zu stellen.',
    '(3) 客户承诺在收到要求后14天内向授权代表和审计师提供所需文件。')

# --- §8 ---
add_heading_styled('§8  Haftung und Freistellung  |  责任与赔偿')

# CHANGED §8(1) — cross-reference to §3(6)
add_changed_clause('(1)',
    'Der Kunde haftet uneingeschränkt für die Richtigkeit, Vollständigkeit und '
    'Rechtzeitigkeit der von ihm übermittelten Daten. Er stellt den Bevollmächtigten '
    'von sämtlichen Ansprüchen Dritter frei, die aus unrichtigen, unvollständigen '
    'oder verspäteten Angaben des Kunden entstehen. § 3 Abs. 6 Sätze 2 bis 5 gelten entsprechend.',
    '(1) 客户对其所提供数据的正确性、完整性和及时性承担无限责任。客户应赔偿授权代表因客户数据错误、'
    '不完整或延迟而产生的所有第三方索赔。第3条第6款第2至第5句相应适用。')

add_clause('(2)',
    'Die Haftung des Bevollmächtigten ist — außer bei Vorsatz, grober '
    'Fahrlässigkeit sowie bei der Verletzung von Leben, Körper oder Gesundheit '
    '— ausgeschlossen.',
    '(2) 除故意、重大过失以及伤害生命、身体或健康外，授权代表的责任被排除。')

add_clause('(3)',
    'Bei einfach fahrlässiger Verletzung von Kardinalpflichten haftet der Bevollmächtigte '
    'begrenzt auf den vertragstypischen, vorhersehbaren Schaden.',
    '(3) 对于因简单过失违反基本义务（即使合同正常履行成为可能的义务），授权代表的责任限于合同典型的、'
    '可预见的损害。')

add_clause('(4)',
    'Die Haftung des Bevollmächtigten ist — außer bei Vorsatz oder grober '
    'Fahrlässigkeit — auf EUR 500,00 pro Einzelfall und EUR 1.000,00 pro '
    'Kalenderjahr (aggregiert) begrenzt.',
    '(4) 除故意或重大过失外，授权代表的总赔偿上限为每案500.00欧元、每个日历年总计1,000.00欧元。')

add_clause('(5)',
    'Der Zahlungsdienstleister FREDDY haftet ausschließlich für eigene vorsätzliche '
    'oder grob fahrlässige Pflichtverletzungen im Rahmen der Zahlungsabwicklung. Eine '
    'Haftung von FREDDY für die fachliche Richtigkeit der Leistungen von LIVANTO ist '
    'ausdrücklich ausgeschlossen.',
    '(5) 支付服务商FREDDY仅对其在支付处理中的故意或重大过失违约负责。FREDDY对LIVANTO服务的'
    '专业性不承担任何责任。')

# --- §9 ---
add_heading_styled('§9  Laufzeit und Kündigung  |  期限与终止')

# CHANGED §9(1) — payment-linked effectiveness
add_changed_clause('(1)',
    'Der Vertrag beginnt am Ersten des auf den Zahlungseingang der ersten '
    'Jahresgrundgebühr folgenden Monats („Vertragsbeginn“). Die erste '
    'Jahresgrundgebühr wird nach Vertragsunterzeichnung in Rechnung gestellt. Auf '
    'Wunsch des Kunden kann der Bevollmächtigte eine sofortige Vertragsgeltung ab '
    'dem Tag des Zahlungseingangs bestätigen. Die anfängliche Laufzeit beträgt '
    '12 Monate ab Vertragsbeginn.',
    '(1) 合同自首笔年度基本费用到账后次月首日起生效（“合同开始”）。首笔年度基本费用'
    '在合同签署后开具发票。应客户要求，授权代表可确认合同自付款到账之日起立即生效。初始期限为'
    '自合同开始起12个月。')

add_clause('(2)',
    'Der Vertrag verlängert sich automatisch um jeweils weitere 12 Monate, sofern er '
    'nicht mit einer Frist von 4 Wochen zum Ende der jeweiligen Vertragslaufzeit '
    'schriftlich gekündigt wird.',
    '(2) 合同自动续期12个月，除非任何一方在合同到期前4周书面通知终止。')

add_clause('(3)',
    'Das Recht zur außerordentlichen Kündigung aus wichtigem Grund bleibt unberührt.',
    '(3) 因重大事由特别终止合同的权利不受影响。')

# CHANGED §9(4) — (d) removed, moved to (4a)
add_changed_clause('(4)',
    'Ein wichtiger Grund für den Bevollmächtigten liegt insbesondere vor bei: '
    '(a) Zahlungsverzug von mehr als 30 Tagen, (b) wiederholter Nichtvorlage der '
    'geschuldeten Verpackungsdaten, (c) Entzug des LUCID-Zugangs durch den Kunden.',
    '(4) 授权代表的重大终止事由特别包括：(a) 逾期付款超过30天，(b) 多次未提交应提交的包装数据，'
    '(c) 客户撤回LUCID访问权限。')

# NEW §9(4a) — undeclared packaging
add_new_clause('(4a)',
    'Bringt der Kunde nicht deklarierte Verpackungsarten ohne vorherige Mitteilung '
    'in Verkehr, steht dem Bevollmächtigten ein sofortiges außerordentliches '
    'Kündigungsrecht zu. Bereits geleistete Jahresgebühren werden in diesem Fall '
    'nicht erstattet. Der Kunde bleibt zur Zahlung etwaiger noch ausstehender dualer '
    'Systemgebühren verpflichtet. Die Freistellungspflicht des Kunden gemäß § 3 Abs. 6 '
    'bleibt unberührt.',
    '(4a) 如客户未经事先通知投放未申报的包装种类，授权代表享有立即特别终止权。此情形下已支付的年费'
    '不予退还。客户仍有义务支付尚未结清的双元系统费用。客户根据第3条第6款的赔偿义务不受影响。')

add_clause('(5)',
    'Nach Vertragsbeendigung ist der Kunde verpflichtet, die Bevollmächtigung '
    'unverzüglich im LUCID-Register zu beenden. Der Bevollmächtigte bestätigt die '
    'Beendigung ebenfalls in LUCID. Versäumt der Kunde die rechtzeitige Löschung, '
    'haftet er für alle daraus entstehenden Schäden.',
    '(5) 合同终止后，客户有义务立即在LUCID登记处终止授权。授权代表同样在LUCID中确认终止。'
    '如客户未及时删除，客户对由此产生的所有损害承担责任。')

add_clause('(6)',
    'Die Pflichten zur Datenmeldung und Abrechnung für das letzte laufende Kalenderjahr '
    'sowie etwaige ausstehende Zahlungsverpflichtungen bestehen über die '
    'Vertragsbeendigung hinaus fort.',
    '(6) 最后一个日历年的数据申报和结算义务以及未清偿的付款义务在合同终止后继续有效。')

# --- §10 ---
add_heading_styled('§10  Datenschutz  |  数据保护')

add_clause('(1)',
    'Der Bevollmächtigte verarbeitet personenbezogene Daten des Kunden ausschließlich '
    'zur Vertragserfüllung und im Einklang mit der Datenschutz-Grundverordnung (DSGVO). '
    'Die vollständige Datenschutzerklärung ist unter [URL] abrufbar.',
    '(1) 授权代表仅在履行合同所需范围内并依据《通用数据保护条例》(GDPR)处理客户的个人数据。'
    '完整的数据保护声明可在[URL]查阅。')

# CHANGED §10(2) — SCC + consent
add_changed_clause('(2)',
    'Die Übermittlung zahlungsbezogener Daten an den Zahlungsdienstleister FREDDY '
    'mit Sitz in der Volksrepublik China erfolgt auf Grundlage der '
    'EU-Standardvertragsklauseln (SCC) gemäß Durchführungsbeschluss (EU) 2021/914 '
    'der Kommission sowie ergänzend auf Grundlage der ausdrücklichen Einwilligung '
    'des Kunden gemäß Art. 49 Abs. 1 lit. a DSGVO. Der Kunde wird hiermit über die '
    'möglichen Risiken einer Datenübermittlung in ein Drittland ohne '
    'Angemessenheitsbeschluss informiert. Die Einwilligung kann jederzeit mit '
    'schriftlicher Mitteilung widerrufen werden. Ein Widerruf führt zur Beendigung '
    'dieses Vertrages mit einer Frist von einem Monat zum Monatsende.',
    '(2) 向位于中华人民共和国的支付服务商FREDDY传输支付相关数据，基于欧盟委员会执行决定'
    '(EU) 2021/914项下的欧盟标准合同条款(SCC)，并补充基于客户根据GDPR第49条第1款(a)项的明确同意。'
    '客户特此被告知向未经充分性认定的第三国传输数据可能存在的风险。该同意可随时通过书面通知撤回。'
    '撤回将导致本合同在月底前一个月通知后终止。')

# --- §11 ---
add_heading_styled('§11  Online-Portal und Kommunikation  |  在线门户与通信')

# CHANGED §11(1) — removed "ausschließlich", added fallback
add_changed_clause('(1)',
    'Vertragsschluss, Datenmeldungen und die Verwaltung der Stammdaten erfolgen '
    'über das Online-Portal des Bevollmächtigten. Bei nachweislichem technischen '
    'Ausfall des Portals kann der Bevollmächtigte die Einreichung per E-Mail gestatten. '
    'Die Zulässigkeit der E-Mail-Einreichung ist vom Bevollmächtigten vorab schriftlich '
    'zu bestätigen.',
    '(1) 合同的签订、数据申报和基本数据管理通过授权代表的在线门户进行。如门户出现可证明的技术故障，'
    '授权代表可允许通过电子邮件提交。电子邮件提交的有效性需由授权代表事先书面确认。')

add_clause('(2)',
    'Mitteilungen an die vom Kunden registrierte E-Mail-Adresse gelten als wirksam '
    'zugestellt. Der Kunde ist verpflichtet, seine Kontaktdaten im Portal stets '
    'aktuell zu halten.',
    '(2) 发送至客户注册邮箱的通知视为有效送达。客户有义务在门户中始终保持其联系数据为最新。')

add_clause('(3)',
    'Der Kunde ist für die Sicherheit seiner Zugangsdaten verantwortlich und '
    'verpflichtet, einen unbefugten Zugriff auf sein Konto unverzüglich zu melden.',
    '(3) 客户对其登录凭证的安全负责，有义务在发现未经授权的账户访问时立即报告。')

# --- §12 ---
add_heading_styled('§12  Schlussbestimmungen  |  最终条款')

add_clause('(1)',
    'Änderungen und Ergänzungen dieses Vertrages bedürfen der Schriftform (§ 126 BGB). '
    'Dies gilt auch für die Aufhebung des Schriftformerfordernisses. Die elektronische '
    'Form (§ 126a BGB) ist der Schriftform gleichgestellt, sofern eine qualifizierte '
    'elektronische Signatur (QES) verwendet wird.',
    '(1) 本合同的修改和补充需采用书面形式（《德国民法典》第126条），书面形式的取消亦需书面形式。'
    '如使用合格电子签名(QES)，电子形式（《德国民法典》第126a条）等同于书面形式。')

add_clause('(2)',
    'Sollte eine Bestimmung dieses Vertrages unwirksam sein oder werden, bleibt die '
    'Wirksamkeit der übrigen Bestimmungen unberührt. Die unwirksame Bestimmung ist durch '
    'eine wirksame zu ersetzen, die dem wirtschaftlichen Zweck der unwirksamen Bestimmung '
    'am nächsten kommt.',
    '(2) 本合同的任何条款如无效或失效，不影响其余条款的效力。无效条款应以最接近其经济目的的有效条款替代。')

add_clause('(3)',
    'Es gilt ausschließlich das Recht der Bundesrepublik Deutschland unter Ausschluss '
    'des UN-Kaufrechts (CISG) und der Kollisionsnormen des internationalen Privatrechts.',
    '(3) 适用法律仅为德意志联邦共和国法律，排除联合国国际货物销售合同公约(CISG)和国际私法的冲突规范。')

add_clause('(4)',
    'Gerichtsstand ist der Sitz des Bevollmächtigten (Dortmund). Der Bevollmächtigte '
    'ist jedoch berechtigt, auch am allgemeinen Gerichtsstand des Kunden zu klagen.',
    '(4) 管辖法院为授权代表所在地（多特蒙德）。但授权代表也有权在客户的普遍管辖法院提起诉讼。')

add_clause('(5)',
    'Dieser Vertrag ist in deutscher und chinesischer Sprache abgefasst. Bei '
    'Auslegungszweifeln oder Widersprüchen zwischen den Sprachfassungen ist die '
    'deutsche Fassung maßgeblich. Die chinesische Übersetzung dient ausschließlich '
    'dem besseren Verständnis.',
    '(5) 本合同以德文和中文两种语言起草。如语言版本之间存在解释疑问或矛盾，以德文版本为准。'
    '中文翻译仅供参考。')

add_page_break()

# ===== ANLAGE A =====
add_heading_styled('Anlage A: Zahlungsabwicklung über FREDDY  |  附件A：FREDDY支付处理')

# NEW FREDDY role description
p = doc.add_paragraph()
r0 = p.add_run('[GEÄNDERT] ')
r0.font.name = 'Arial'
r0.font.size = Pt(8)
r0.font.color.rgb = RED
r0.bold = True
r1 = p.add_run(
    'FREDDY (Shanghai) Information Consulting Ltd., Niederlassung Huainan '
    '(nachfolgend „FREDDY“), unterstützt LIVANTO bei der Zahlungsabwicklung, '
    'Markterschließung und Kundenkommunikation im chinesischsprachigen Raum. '
    'Die Parteien vereinbaren folgende Zahlungsabwicklung:'
)
r1.font.name = 'Arial'
r1.font.size = Pt(10)

p2 = doc.add_paragraph()
r2a = p2.add_run('[已修改] ')
r2a.font.name = 'Microsoft YaHei'
r2a.font.size = Pt(7)
r2a.font.color.rgb = RED
r2a.bold = True
r2b = p2.add_run(
    '福瑞笛（上海）信息咨询有限公司淮南分公司（以下称“FREDDY”）在中文地区为LIVANTO'
    '提供支付处理、市场拓展和客户沟通支持。双方约定如下支付处理方式：'
)
r2b.font.name = 'Microsoft YaHei'
r2b.font.size = Pt(9)
r2b.font.color.rgb = GRAY

add_clause('(1)',
    'FREDDY wird als Zahlungsdienstleister eingeschaltet. Die Zahlung an FREDDY '
    'gilt als wirksame und schuldbefreiende Zahlung des Kunden an LIVANTO.',
    '(1) FREDDY被指定为支付服务商。客户向FREDDY的付款视为向LIVANTO有效且解除债务的付款。')

add_clause('(2)',
    'FREDDY leitet die vereinnahmten Beträge nach Abzug einer vertraglich vereinbarten '
    'Servicegebühr innerhalb von 14 Tagen nach schriftlicher Anforderung durch LIVANTO '
    'an LIVANTO weiter.',
    '(2) FREDDY在收到LIVANTO书面要求后14天内，将扣除约定服务费后的款项转交LIVANTO。')

add_clause('(3)',
    'LIVANTO leitet die dualen Systemgebühren an das jeweils beauftragte duale System weiter.',
    '(3) LIVANTO将双元系统费用转交给相应的双元系统。')

add_clause('(4)',
    'Akzeptierte Zahlungsmethoden: WeChat Pay, Alipay, Banküberweisung. FREDDY kann '
    'die Liste der akzeptierten Zahlungsmethoden jederzeit aktualisieren.',
    '(4) 接受的支付方式：微信支付、支付宝、银行转账。FREDDY可随时更新接受的支付方式列表。')

add_clause('(5)',
    'FREDDY haftet ausschließlich für eigene vorsätzliche oder grob fahrlässige '
    'Pflichtverletzungen im Rahmen der Zahlungsabwicklung.',
    '(5) FREDDY仅对其在支付处理中的故意或重大过失违约负责。')

# NEW catch-all clause
add_new_clause('(6)',
    'FREDDY handelt ausschließlich als Zahlungsdienstleister und Marktkooperationspartner '
    'von LIVANTO. Diese Anlage A begründet keine eigenen Rechte von FREDDY auf Kündigung, '
    'Vertragsänderung oder Befreiung von Vertragspflichten der Parteien. FREDDY ist nicht '
    'Vertragspartei des Hauptvertrages.',
    '(6) FREDDY仅作为LIVANTO的支付服务商和市场合作方行事。本附件A不赋予FREDDY任何合同解除权、'
    '合同修改权或合同义务豁免权。FREDDY不是主合同的合同当事人。')

# Renumbered as (7)
add_clause('(7)',
    'Im Falle eines Widerspruchs zwischen dieser Anlage A und dem Hauptvertrag gehen '
    'die Regelungen dieser Anlage A vor.',
    '(7) 如本附件A与主合同有冲突，以本附件A的规定为准。')

add_page_break()

# ===== ANLAGE B =====
add_heading_styled('Anlage B: Verpackungsdeklaration des Kunden  |  附件B：客户包装申报表')

add_de(
    'Der Kunde erklärt bei Vertragsschluss nach bestem Wissen folgende '
    'Verpackungsarten und -mengen im Sinne des Verpackungsgesetzes. Diese Angaben '
    'dienen als Grundlage für den Abschluss des dualen Systemvertrages.',
    '客户在合同签订时据其所知声明以下《包装法》意义上的包装种类和数量。该声明作为签订双元系统合同的依据。'
)

# Packaging declaration table
table = doc.add_table(rows=9, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Materialfraktion\n材料类别', 'Kategorie\n类别', 'Geschätzte Menge\n预估年量(kg)', 'Produktbeispiel\n产品举例']
for col, h in enumerate(headers):
    cell = table.cell(0, col)
    set_cell_shading(cell, HEADER_BG)
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.font.name = 'Arial'
    run.font.size = Pt(8)
    run.bold = True

materials = [
    'Glas / 玻璃', 'Papier/Pappe/Karton / 纸板', 'Eisenmetalle / 黑色金属',
    'Aluminium / 铝', 'Kunststoffe / 塑料', 'Verbunde / 复合材料',
    'Getränkekarton / 饮料纸盒', 'Sonstige / 其他'
]
for row_idx, mat in enumerate(materials):
    cell = table.cell(row_idx + 1, 0)
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(mat)
    run.font.name = 'Arial'
    run.font.size = Pt(8)
    for col in range(1, 4):
        table.cell(row_idx + 1, col).text = ''

doc.add_paragraph('')
add_de(
    'Der Kunde bestätigt die Richtigkeit der vorstehenden Angaben nach bestem Wissen und Gewissen.',
    '客户据其所知确认上述信息真实准确。'
)

add_page_break()

# ===== ANLAGE C =====
add_heading_styled('Anlage C: Gebührenordnung  |  附件C：费用表')

add_de(
    'Gültig ab Vertragsbeginn. Alle Beträge verstehen sich netto zuzüglich der '
    'gesetzlichen Mehrwertsteuer.',
    '自合同开始生效。所有金额为净价，另加法定增值税。'
)

# Fee table
table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

for col, h in enumerate(['Servicestufe\n服务等级', 'Jahresgrundgebühr\n年度基本费', 'Leistungsumfang / 服务范围']):
    cell = table.cell(0, col)
    set_cell_shading(cell, HEADER_BG)
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.font.name = 'Arial'
    run.font.size = Pt(8)
    run.bold = True

fee_data = [
    ('Basis\n基础', 'EUR 89,00',
     'Kern-AR-Dienstleistungen: Systembeteiligung, LUCID-Datenmeldung, Behördenkorrespondenz, '
     'gesetzliche Informationspflichten\n核心AR服务：系统参与、LUCID数据申报、官方通信、法定信息义务'),
    ('Standard\n标准', 'EUR 159,00',
     'Basis + prioritäre Reaktionszeit (48 h) + erweiterte Klassifizierungsberatung + '
     'jährlicher Compliance-Kurzbericht\n基础+优先响应时间(48小时)+扩展分类咨询+年度合规简报'),
    ('Premium\n高级', 'EUR 249,00',
     'Standard + Vollständigkeitserklärungs-Koordination + persönlicher Ansprechpartner + '
     '24h-Reaktionszeit\n标准+完整性声明协调+专属客户经理+24小时响应时间'),
]

for row_idx, (level, price, desc) in enumerate(fee_data):
    for col_idx, text in enumerate([level, price, desc]):
        cell = table.cell(row_idx + 1, col_idx)
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(8)
        if col_idx == 0:
            run.bold = True

doc.add_paragraph('')
add_simple('Zusätzliche Leistungen:  |  附加服务：', bold=True)
add_simple('  •  ZSVR-Klassifizierungsantrag: EUR 150,00  |  ZSVR分类申请：150.00欧元')
add_simple('  •  Vollständigkeitserklärung (Prüfer + Bearbeitung): Prüferkosten + EUR 50,00  |  完整性声明（审计+手续费）：审计费用 + 50.00欧元')
add_simple('  •  Mahnung bei Zahlungsverzug: EUR 15,00  |  逾期催款：15.00欧元')
add_simple('  •  Duale Systemgebühren: Weitergabe zu Selbstkosten (ohne Aufschlag)  |  双元系统费用：按成本价过账（不加价）')

add_page_break()

# ===== SIGNATURES =====
add_heading_styled('Unterschriften  |  签署')

doc.add_paragraph('')
add_signature_table()

doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Bevollmächtigungsvertrag V3.0 — LIVANTO GmbH — Stand: Juni 2026')
run.font.name = 'Arial'
run.font.size = Pt(7.5)
run.font.color.rgb = LIGHT_GRAY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    'Änderungen ggü. V2: §3(4), §3(6), §5(2), §5(5), §6(5), §6(6), §8(1), '
    '§9(1), §9(4a neu), §10(2), §11(1), Anlage A'
)
run.font.name = 'Arial'
run.font.size = Pt(6.5)
run.font.color.rgb = LIGHT_GRAY


# ===== SAVE =====
output_path = r'C:\Users\ThinkPad\Desktop\AI-Projects\Freddy-Website\projects\contracts\LIVANTO\Bevollmächtigungsvertrag_03.docx'
doc.save(output_path)
print(f'DOCX generated successfully: {output_path}')
print(f'Size: {os.path.getsize(output_path) / 1024:.1f} KB')
